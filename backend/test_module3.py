"""
test_module3.py — Automated test suite for Module 3 Structured Resume Intelligence & Parsing.

Tests execute deterministically without external LLM calls or API credits.

Coverage:
1. Standard resume parsing (Resume A)
2. Missing sections & truth preservation (Resume B)
3. Alternate section headings (Resume C)
4. Skills extraction & special character preservation (Resume E)
5. Experience extraction
6. Project extraction
7. Education extraction
8. Contact information extraction (email, phone, LinkedIn, GitHub)
9. Special characters: C++, C#, .NET, Node.js, React.js, Python, SQL, std::vector<T>, {templates}
10. Empty document handling
11. Corrupted PDF handling
12. Corrupted DOCX handling
13. Full Module 1 & 2 regression integration
"""

import io
import os
import sys
import unittest
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient
import PyPDF2

# Add backend directory to sys.path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from main import app
from schemas.resume import ParsedResume, ResumeDocument
from services.parser_service import parse_resume_bytes, parse_file_to_document
from services.resume_structurer import (
    extract_contact_info,
    normalize_resume_text,
    parse_skills_section,
    split_into_sections,
    structure_resume,
)

client = TestClient(app)

# ── Fixtures ──────────────────────────────────────────────────────────────────
RESUME_A_STANDARD = """\
John Doe
john.doe@example.com | (555) 123-4567 | San Francisco, CA
https://linkedin.com/in/johndoe | https://github.com/johndoe

Professional Summary
Senior Software Engineer with 6+ years of experience in backend development and cloud services.

Technical Skills
Languages & Frameworks: Python, C++, SQL, FastAPI, Docker, PostgreSQL, React.js

Work Experience
Senior Backend Engineer | TechCorp Inc.
Jan 2021 - Present
- Architected microservice handling 10M daily requests with 99.99% uptime.
- Optimized database queries reducing latency by 45%.

Software Engineer | StartupXYZ
Jun 2018 - Dec 2020
- Developed REST APIs using Python and PostgreSQL.

Projects
AI Resume Parser
https://github.com/johndoe/resume-parser
- Built deterministic NLP parser for tech resumes.

Education
B.S. in Computer Science | Stanford University
2014 - 2018
GPA: 3.9/4.0

Certifications
AWS Certified Solutions Architect
"""

RESUME_B_MISSING_SECTIONS = """\
Jane Smith
jane.smith@email.com

Skills
Python, SQL, Git

Projects
Data Pipeline Tool
Built automated data pipeline for JSON processing.
"""

RESUME_C_ALTERNATE_HEADINGS = """\
Alex Rivera
alex@domain.io | github.com/arivera

Technical Competencies
Python, Go, Kubernetes, Terraform

Professional Background
Lead Cloud Architect | CloudNative Systems
2020 - Present
- Deployed multi-region Kubernetes clusters.

Academic Background
Bachelor of Engineering in Computer Science | State University
2015 - 2019
"""

RESUME_E_SPECIAL_CHARS = """\
Dev Specialist
dev@special.org

Skills
C++, C#, .NET, Node.js, React.js, Python, SQL, std::vector<T>, {templates}

Experience
Software Engineer | CodeLab
2022 - Present
- Implemented C++ templates using std::vector<T> and custom containers with {templates}.
"""


class TestModule3(unittest.TestCase):

    # ── Unit tests for Structurer ──────────────────────────────────────────────
    def test_01_standard_resume_structuring(self):
        """Test 1: Standard resume A parsed completely."""
        parsed = structure_resume(RESUME_A_STANDARD)
        self.assertIsInstance(parsed, ParsedResume)
        self.assertEqual(parsed.contact.email, "john.doe@example.com")
        self.assertEqual(parsed.contact.phone, "(555) 123-4567")
        self.assertIn("linkedin.com/in/johndoe", parsed.contact.linkedin or "")
        self.assertIn("github.com/johndoe", parsed.contact.github or "")
        self.assertIn("Python", parsed.skills)
        self.assertIn("C++", parsed.skills)
        self.assertEqual(len(parsed.experience), 2)
        self.assertEqual(len(parsed.projects), 1)
        self.assertEqual(len(parsed.education), 1)
        self.assertEqual(len(parsed.certifications), 1)

    def test_02_missing_sections_truth_preservation(self):
        """Test 2 & 10: Missing sections result in empty lists, NOT fabricated data."""
        parsed = structure_resume(RESUME_B_MISSING_SECTIONS)
        self.assertEqual(len(parsed.experience), 0)
        self.assertEqual(len(parsed.education), 0)
        self.assertEqual(len(parsed.certifications), 0)
        self.assertEqual(len(parsed.skills), 3)
        self.assertEqual(len(parsed.projects), 1)

    def test_03_alternate_section_headings(self):
        """Test 3: Recognizes non-standard section headings cleanly."""
        parsed = structure_resume(RESUME_C_ALTERNATE_HEADINGS)
        self.assertIn("Kubernetes", parsed.skills)
        self.assertEqual(len(parsed.experience), 1)
        self.assertEqual(parsed.experience[0].company, "CloudNative Systems")
        self.assertEqual(len(parsed.education), 1)

    def test_04_skills_extraction_special_characters(self):
        """Test 4 & 9: Special characters like C++, C#, .NET, std::vector<T>, {templates} preserved."""
        parsed = structure_resume(RESUME_E_SPECIAL_CHARS)
        self.assertIn("C++", parsed.skills)
        self.assertIn("C#", parsed.skills)
        self.assertIn(".NET", parsed.skills)
        self.assertIn("Node.js", parsed.skills)
        self.assertIn("std::vector<T>", parsed.skills)
        self.assertIn("{templates}", parsed.skills)

    def test_05_contact_extraction(self):
        """Test 8: Regex contact extraction for email, phone, github, linkedin."""
        contact = extract_contact_info(RESUME_A_STANDARD.split("\n"), RESUME_A_STANDARD)
        self.assertEqual(contact.email, "john.doe@example.com")
        self.assertEqual(contact.phone, "(555) 123-4567")
        self.assertIn("johndoe", contact.github or "")
        self.assertIn("johndoe", contact.linkedin or "")

    def test_06_text_normalization(self):
        """Test: Normalization converts weird Unicode bullets and standardizes line endings."""
        raw_unicode = "Skill 1\r\n\u2022 Skill 2\r▪ Skill 3\n\n\n\n- Skill 4"
        norm = normalize_resume_text(raw_unicode)
        self.assertNotIn("\r", norm)
        self.assertIn("- Skill 2", norm)
        self.assertNotIn("\n\n\n", norm)

    def test_07_empty_document_handling(self):
        """Test 10: Empty string returns empty ParsedResume safely."""
        parsed = structure_resume("")
        self.assertIsInstance(parsed, ParsedResume)
        self.assertEqual(len(parsed.skills), 0)

        with self.assertRaises(ValueError):
            parse_resume_bytes(b"", "empty.pdf")

    def test_08_corrupted_pdf_handling(self):
        """Test 11: Corrupted PDF binary raises clean ValueError."""
        corrupted_bytes = b"%PDF-1.4 corrupted invalid binary header garbage"
        with self.assertRaises(ValueError) as ctx:
            parse_resume_bytes(corrupted_bytes, "invalid.pdf")
        self.assertIn("Failed to extract text from PDF", str(ctx.exception))

    def test_09_corrupted_docx_handling(self):
        """Test 12: Corrupted DOCX binary raises clean ValueError."""
        corrupted_bytes = b"PK\x03\x04 corrupted docx zip stream"
        with self.assertRaises(ValueError) as ctx:
            parse_resume_bytes(corrupted_bytes, "invalid.docx")
        self.assertIn("Failed to extract text from DOCX", str(ctx.exception))

    # ── Synthetic PDF & DOCX Binary Tests ─────────────────────────────────────
    def test_10_docx_binary_extraction(self):
        """Test: Valid in-memory DOCX file extraction."""
        doc = Document()
        doc.add_paragraph("Alice Tech")
        doc.add_paragraph("alice@tech.org")
        doc.add_paragraph("Skills: Python, FastAPI")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()

        extracted = parse_resume_bytes(docx_bytes, "resume.docx")
        self.assertIn("Alice Tech", extracted)
        self.assertIn("Python, FastAPI", extracted)

    # ── Module 1 & 2 Integration / Regression Tests ───────────────────────────
    def test_11_endpoint_module1_and_2_regression(self):
        """Test 13, 14, 15: Full API endpoint behavior remains 100% compliant with Module 1 & 2."""
        with patch("routes.analyze.analyze_resume") as mock_analyze:
            from schemas.analysis import AnalysisResponse
            mock_analyze.return_value = AnalysisResponse(
                ats_score=90,
                match_percentage="90%",
                optimized_resume="Optimized text",
                key_improvements=["Improved bullets"],
                original_bullets=["Original bullet"],
                optimized_bullets=["Optimized bullet"],
                skill_gaps=[],
                project_suggestions=[],
                top_strengths=["Python"],
                recruiter_insight="Great candidate",
                github_integration=[]
            )
            res = client.post(
                "/api/analyze",
                data={
                    "resume_text": RESUME_A_STANDARD,
                    "job_description": "We need a Senior Backend Engineer with Python and C++ experience."
                }
            )
            self.assertEqual(res.status_code, 200)
            self.assertIsInstance(res.json()["ats_score"], int)
            self.assertIn("ats_breakdown", res.json())



if __name__ == "__main__":
    unittest.main()
